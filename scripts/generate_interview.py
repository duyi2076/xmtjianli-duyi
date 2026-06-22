"""Generate the interview story docx."""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from pathlib import Path


def _set_chinese_font(run, font_name='SimSun', size=10.5):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)


def _add_heading(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    _set_chinese_font(run, 'Microsoft YaHei', size)
    p.space_after = Pt(8)
    return p


def _add_para(doc, text, size=10.5, bold=False):
    last = None
    text = str(text or "").replace("\\n", "\n")
    for part in text.split("\n"):
        p = doc.add_paragraph()
        if part:
            run = p.add_run(part)
            run.bold = bold
            _set_chinese_font(run, 'SimSun', size)
        p.space_after = Pt(6)
        last = p
    return last


def generate_interview(data, output_path):
    """
    Generate interview story docx.

    Args:
        data: dict with keys:
            - name: str
            - job_intent: str
            - oral_version: str
            - detailed_version: str
            - ai_notes: str
        output_path: path to save the generated docx
    """
    output_path = Path(output_path).expanduser()

    doc = Document()

    _add_heading(doc, data.get("name", ""), 16, True)
    _add_heading(doc, f"{data.get('job_intent', '')} 面试辅助材料", 12, False)
    _add_para(doc, "")

    _add_heading(doc, "版本一：面试口述版", 12, True)
    _add_para(doc, "─" * 40)
    _add_para(doc, data.get("oral_version", ""), 10.5)
    _add_para(doc, "")

    _add_heading(doc, "版本二：个人经历详细版", 12, True)
    _add_para(doc, "─" * 40)
    _add_para(doc, data.get("detailed_version", ""), 10.5)
    _add_para(doc, "")

    _add_heading(doc, "AI 补充说明", 12, True)
    _add_para(doc, "─" * 40)
    _add_para(doc, data.get("ai_notes", ""), 10.5)

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    import json
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument("--data", required=True, help="Path to JSON data file")
    parser.add_argument("--output", required=True, help="Output docx path")
    args = parser.parse_args()

    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)

    generate_interview(data, args.output)
    print(f"Interview doc saved to: {args.output}")

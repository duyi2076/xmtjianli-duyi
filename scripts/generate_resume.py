"""Generate the corrected resume docx from template."""
import base64
import os
import re
import tempfile
import zipfile
from docx import Document
from docx.oxml import OxmlElement
from pathlib import Path


def _find_para_index(doc, predicate):
    for i, p in enumerate(doc.paragraphs):
        if predicate(p.text):
            return i
    return -1


def _delete_paragraphs(doc, start_idx, count):
    for _ in range(count):
        if start_idx < len(doc.paragraphs):
            p = doc.paragraphs[start_idx]._element
            p.getparent().remove(p)


def _remove_duplicate_heading(doc, text):
    seen = False
    for para in list(doc.paragraphs):
        if para.text.strip() != text:
            continue
        if not seen:
            seen = True
            continue
        p = para._element
        p.getparent().remove(p)


def _remove_empty_project_blocks(doc):
    i = 0
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        is_empty_project = re.fullmatch(r"项目 \d+：", text)
        following = [p.text.strip() for p in doc.paragraphs[i + 1:i + 4]]
        is_empty_body = following == ["负责事项：", "项目介绍：", "取得成果："]
        if is_empty_project and is_empty_body:
            start = i
            if i > 0 and not doc.paragraphs[i - 1].text.strip():
                start = i - 1
            _delete_paragraphs(doc, start, i + 4 - start)
            i = max(start - 1, 0)
            continue
        i += 1


def _insert_paragraph_after(ref, text):
    new_p = OxmlElement('w:p')
    if hasattr(ref, '_element'):
        ref._element.addnext(new_p)
    else:
        ref.addnext(new_p)

    new_run = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_run.append(new_t)
    new_p.append(new_run)
    return new_p


def _replace_placeholders(doc, replacements):
    for para in doc.paragraphs:
        for run in para.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for old, new in replacements.items():
                            if old in run.text:
                                run.text = run.text.replace(old, new)


def _ensure_template(template_path):
    """Return a usable template path. If the .docx is missing, decode from base64 asset."""
    template_path = Path(template_path).expanduser()
    if template_path.exists():
        return template_path

    base64_path = template_path.parent / "template_docx_base64.txt"
    if not base64_path.exists():
        raise FileNotFoundError(f"找不到简历模板：{template_path}，也找不到 base64 编码文件：{base64_path}")

    decoded = template_path.parent / "template.docx"
    with open(base64_path, "r", encoding="utf-8") as f:
        data = base64.b64decode(f.read().strip())
    with open(decoded, "wb") as f:
        f.write(data)
    return decoded


def _strip_template_textboxes(docx_path):
    """Remove template drawing text boxes that python-docx cannot replace reliably."""
    docx_path = Path(docx_path)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()

    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp.name, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                xml = data.decode("utf-8")
                xml = re.sub(r"<w:drawing>.*?</w:drawing>", "", xml, flags=re.DOTALL)
                xml = re.sub(r"<w:pict>.*?</w:pict>", "", xml, flags=re.DOTALL)
                data = xml.encode("utf-8")
            zout.writestr(item, data)

    os.replace(tmp.name, docx_path)
    os.chmod(docx_path, 0o644)


def generate_resume(data, output_path, template_path=None):
    """
    Generate resume docx.

    Args:
        data: dict with keys:
            - personal_info: dict with name, phone, email, age, city
            - advantages: list of 5 strings
            - experiences: list of dicts with company, position, period, project_name, duty, intro, result
            - skills: dict with professional, awards, other
            - education: dict with school, major, period, major_courses, elective_courses
        output_path: path to save the generated docx
        template_path: optional path to template.docx
    """
    if template_path is None:
        template_path = Path(__file__).parent.parent / "assets" / "template.docx"

    template_path = _ensure_template(template_path)

    output_path = Path(output_path).expanduser()

    # Build placeholder replacements
    info = data["personal_info"]
    advantages = data["advantages"]
    experiences = data["experiences"]
    skills = data["skills"]
    education = data["education"]
    job_intent = data.get("job_intent") or info.get("job_intent") or "求职意向待确认"
    advantage_labels = data.get(
        "advantage_labels",
        ["文案撰写能力", "平台运营能力", "视频拍摄剪辑能力", "直播运营能力", "附加能力"],
    )

    replacements = {
        "【姓名】": info.get("name", ""),
        "电话/邮箱/年龄/现居地：待补充": (
            f"电话：{info.get('phone', '待补充')} | "
            f"邮箱：{info.get('email', '待补充')} | "
            f"年龄：{info.get('age', '待补充')} | "
            f"现居地：{info.get('city', '待补充')}"
        ),
        "【专业技能】": skills.get("professional", "待补充"),
        "【证书奖项】": skills.get("awards", "待补充"),
        "【其他证书】": skills.get("other", "待补充"),
        "【学校名称】": education.get("school", ""),
        "【专业】": education.get("major", ""),
        "【时间段】": education.get("period", ""),
        "【主修课程】": education.get("major_courses", ""),
        "【选修课程】": education.get("elective_courses", ""),
        "文案撰写能力": advantage_labels[0],
        "平台运营能力": advantage_labels[1],
        "视频拍摄剪辑能力": advantage_labels[2],
        "直播运营能力": advantage_labels[3],
        "【附加能力】": advantage_labels[4],
    }

    for i, adv in enumerate(advantages[:5], 1):
        label = advantage_labels[i - 1]
        adv = str(adv or "").strip()
        # Avoid duplicated label if adv already starts with the label
        if adv.startswith(f"{label}：") or adv.startswith(f"{label}:"):
            adv = adv[len(label) + 1:].strip()
        replacements[f"【个人优势{i}】"] = adv

    # Clear project 2 placeholders under company 1
    for key in ["【项目名称2】", "【负责事项2】", "【项目介绍2】", "【项目结果2】"]:
        replacements[key] = ""

    # Map experiences to template placeholders
    # Template: company1 has project1 and project2; company2 has project3; company3 project4; company4 project5
    # We want: each experience has its own company and one project
    # So company1=exp[0], project1=exp[0]
    # company2=exp[1], project3=exp[1]
    # company3=exp[2], project4=exp[2]
    # company4=exp[3], project5=exp[3]
    # company5=exp[4], project5 is appended after project4 result

    company_project_map = [
        (1, 1),  # company1 -> project1
        (2, 3),  # company2 -> project3 (template project number)
        (3, 4),  # company3 -> project4
        (4, 5),  # company4 -> project5
    ]

    for company_idx, project_template_idx in company_project_map:
        exp = experiences[company_idx - 1] if company_idx <= len(experiences) else {}
        replacements[f"【公司{company_idx}】"] = exp.get("company", "")
        replacements[f"【职位{company_idx}】"] = exp.get("position", "")
        replacements[f"【时间段{company_idx}】"] = exp.get("period", "")
        replacements[f"【项目名称{project_template_idx}】"] = exp.get("project_name", "")
        replacements[f"【负责事项{project_template_idx}】"] = exp.get("duty", "")
        replacements[f"【项目介绍{project_template_idx}】"] = exp.get("intro", "")
        replacements[f"【项目结果{project_template_idx}】"] = exp.get("result", "")

    # Experience 5 will be appended after project4 result
    extra_experience = experiences[4] if len(experiences) >= 5 else None

    doc = Document(template_path)
    _replace_placeholders(doc, replacements)
    _remove_duplicate_heading(doc, "核心优势")

    first_para = doc.paragraphs[0]
    contact = (
        f"电话：{info.get('phone', '待补充')} | "
        f"邮箱：{info.get('email', '待补充')} | "
        f"年龄：{info.get('age', '待补充')} | "
        f"现居地：{info.get('city', '待补充')}"
    )
    first_para.insert_paragraph_before(info.get("name", "姓名待补充"))
    first_para.insert_paragraph_before(f"求职意向：{job_intent}")
    first_para.insert_paragraph_before(contact)

    # Delete template project 2 (4 paragraphs)
    idx_p2 = _find_para_index(doc, lambda t: "项目 2：" in t)
    if idx_p2 >= 0:
        _delete_paragraphs(doc, idx_p2, 4)

    # Rename project numbers: project3->2, project4->3, project5->4
    for old_prefix, new_prefix in [("项目 3：", "项目 2："), ("项目 4：", "项目 3："), ("项目 5：", "项目 4：")]:
        idx = _find_para_index(doc, lambda t, op=old_prefix: op in t)
        if idx >= 0:
            para = doc.paragraphs[idx]
            for run in para.runs:
                if old_prefix in run.text:
                    run.text = run.text.replace(old_prefix, new_prefix)

    # Append extra experience 5 after project 4 result
    if extra_experience:
        idx_p4 = _find_para_index(doc, lambda t: "项目 4：" in t)
        if idx_p4 >= 0:
            idx_p4_result = idx_p4 + 3
            if idx_p4_result < len(doc.paragraphs):
                ref = doc.paragraphs[idx_p4_result]._element
                ref = _insert_paragraph_after(
                    ref,
                    f"{extra_experience.get('company', '')}                    "
                    f"{extra_experience.get('position', '')}                                  "
                    f"{extra_experience.get('period', '')}"
                )
                ref = _insert_paragraph_after(ref, f"项目 5：{extra_experience.get('project_name', '')}")
                ref = _insert_paragraph_after(ref, f"负责事项：{extra_experience.get('duty', '')}")
                ref = _insert_paragraph_after(ref, f"项目介绍：{extra_experience.get('intro', '')}")
                _insert_paragraph_after(ref, f"取得成果：{extra_experience.get('result', '')}")

    _remove_empty_project_blocks(doc)
    doc.save(output_path)
    _strip_template_textboxes(output_path)
    return output_path


if __name__ == "__main__":
    import json
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument("--data", required=True, help="Path to JSON data file")
    parser.add_argument("--output", required=True, help="Output docx path")
    parser.add_argument("--template", help="Optional template docx path")
    args = parser.parse_args()

    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)

    generate_resume(data, args.output, args.template)
    print(f"Resume saved to: {args.output}")
